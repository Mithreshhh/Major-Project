"""
skill_extraction.py

Extracts skill/topic phrases from raw syllabus text using spaCy.

Pipeline:
    1. Get raw text — either passed in directly, or pulled from a PDF
       (PyMuPDF) or DOCX (python-docx) file.
    2. Run spaCy over the text and collect candidate phrases from noun
       chunks and a small set of relevant named-entity labels.
    3. Clean each candidate: strip filler lead-ins ("fundamentals of",
       "introduction to", ...), drop stopword-only / punctuation-only /
       numeric-only phrases, normalize whitespace, and dedupe.

Run this file directly to execute a self-test against a sample syllabus
paragraph:
    python -m app.skill_extraction

Requires the spaCy small English model:
    python -m spacy download en_core_web_sm
"""

import re
from pathlib import Path

import spacy

# PyMuPDF and python-docx are optional at import time so that pure-text
# extraction (extract_skills_from_text) still works even if a consumer
# hasn't installed the file-parsing extras.
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx  # python-docx
except ImportError:
    docx = None


# ---------------------------------------------------------------------------
# spaCy model (loaded lazily and cached, since loading is relatively slow)
# ---------------------------------------------------------------------------

_NLP = None


def get_nlp():
    """Load and cache the spaCy pipeline used for extraction."""
    global _NLP
    if _NLP is None:
        try:
            _NLP = spacy.load("en_core_web_sm")
        except OSError as exc:
            raise RuntimeError(
                'spaCy model "en_core_web_sm" is not installed. Install it with:\n'
                "    python -m spacy download en_core_web_sm"
            ) from exc
    return _NLP


# ---------------------------------------------------------------------------
# File parsing (PDF / DOCX -> raw text)
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_path) -> str:
    """Extract raw text from a PDF file using PyMuPDF."""
    if fitz is None:
        raise ImportError("PyMuPDF is not installed. Run: pip install pymupdf")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    try:
        with fitz.open(path) as pdf:
            return "\n".join(page.get_text() for page in pdf)
    except Exception as exc:  # noqa: BLE001 - surface as a clear, single error type
        raise RuntimeError(f"Failed to parse PDF {path}: {exc}") from exc


def extract_text_from_docx(file_path) -> str:
    """Extract raw text from a DOCX file using python-docx (paragraphs + tables)."""
    if docx is None:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        document = docx.Document(path)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Failed to parse DOCX {path}: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Syllabi often list topics/units in tables rather than paragraphs.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def extract_text_from_file(file_path) -> str:
    """Dispatch to the right parser based on file extension (.pdf / .docx)."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    raise ValueError(f"Unsupported file type '{suffix}' (expected .pdf or .docx)")


# ---------------------------------------------------------------------------
# Noise removal helpers
# ---------------------------------------------------------------------------

# Lead-in phrases that add no topical meaning ("fundamentals of X" -> "X").
FILLER_PHRASES = [
    "fundamentals of",
    "fundamental concepts of",
    "introduction to",
    "an introduction to",
    "basics of",
    "basic concepts of",
    "overview of",
    "principles of",
    "concepts of",
    "elements of",
    "essentials of",
    "foundations of",
    "understanding of",
    "study of",
    "applications of",
]

_FILLER_PATTERN = re.compile(
    r"^\s*(?:" + "|".join(re.escape(p) for p in FILLER_PHRASES) + r")\s+",
    re.IGNORECASE,
)

# Named-entity labels worth keeping as skill candidates (technologies, products,
# languages, organizations/frameworks). Other labels (DATE, GPE, CARDINAL, ...)
# are not skill-relevant and are ignored.
RELEVANT_ENTITY_LABELS = {"ORG", "PRODUCT", "LANGUAGE"}

# Generic academic/course vocabulary that isn't a skill/topic on its own
# (e.g. "the fundamentals", "this course", "strong understanding"). Checked
# against the whole phrase and against its last word, since these terms
# mostly act as the (meaningless) head noun of throwaway phrases.
GENERIC_NOISE_WORDS = {
    "introduction", "fundamentals", "fundamental", "basics", "basic",
    "overview", "principles", "principle", "concepts", "concept",
    "elements", "element", "essentials", "essential", "foundations",
    "foundation", "understanding", "study", "studies", "applications",
    "application", "course", "courses", "student", "students",
    "semester", "semesters", "end", "class", "classes", "unit", "units",
    "module", "modules", "chapter", "chapters", "lecture", "lectures",
    "lab", "labs", "assignment", "assignments", "week", "weeks",
}

MIN_PHRASE_CHARS = 2
MAX_PHRASE_WORDS = 6


def _strip_filler(phrase: str) -> str:
    """Repeatedly strip leading filler phrases (handles stacked cases)."""
    previous = None
    cleaned = phrase
    while previous != cleaned:
        previous = cleaned
        cleaned = _FILLER_PATTERN.sub("", cleaned).strip()
    return cleaned


def _normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _strip_edge_punct(text: str) -> str:
    return text.strip(" .,;:-–—()[]{}\"'")


def _is_noise_span(span) -> bool:
    """A span is noise if it's empty, purely numeric/punctuation, or made
    entirely of stopwords (e.g. "the", "which is used") once trimmed."""
    text = span.text.strip()
    if not text:
        return True
    if not any(char.isalpha() for char in text):
        return True
    if all(token.is_stop or token.is_punct or token.is_space for token in span):
        return True
    return False


def _is_generic_noise(phrase: str) -> bool:
    """Drop throwaway academic phrases like "the fundamentals" or "strong
    understanding" whose meaning carries no topic information. Only checked
    against the last word for short (<=2 word) phrases, since longer ones
    ("software engineering principles") use the same words as real modifiers."""
    words = phrase.lower().split()
    if not words:
        return True
    if phrase.lower() in GENERIC_NOISE_WORDS:
        return True
    return len(words) <= 2 and words[-1] in GENERIC_NOISE_WORDS


def _drop_leading_det_or_pron(span, doc):
    """Trim leading determiners/pronouns spaCy sometimes folds into noun
    chunks (e.g. "the database" -> "database")."""
    start = span.start
    while start < span.end and doc[start].pos_ in ("DET", "PRON"):
        start += 1
    return doc[start:span.end]


# ---------------------------------------------------------------------------
# Core extraction
# ---------------------------------------------------------------------------

def extract_skills_from_text(text: str) -> list:
    """
    Extract a clean, deduplicated list of skill/topic phrases from raw text.

    Returns an empty list for empty/whitespace-only input rather than raising,
    since "no skills found" is a valid, expected outcome.
    """
    if not text or not text.strip():
        return []

    # Collapse newlines/tabs to single spaces first. PDF/DOCX extraction
    # frequently breaks phrases across lines (e.g. "the\nfundamentals"),
    # which otherwise confuses POS-based trimming and phrase boundaries.
    text = _normalize_whitespace(text)

    nlp = get_nlp()
    doc = nlp(text)

    candidates = []

    # Noun chunks capture most curriculum phrasing, e.g. "object-oriented
    # programming", "relational database design". spaCy occasionally folds
    # a comma-joined list into a single chunk, so split those back apart.
    for chunk in doc.noun_chunks:
        span = _drop_leading_det_or_pron(chunk, doc)
        if len(span) == 0 or _is_noise_span(span):
            continue
        candidates.extend(part.strip() for part in span.text.split(",") if part.strip())

    # Named entities catch specific technologies/products that don't always
    # form a full noun chunk on their own (e.g. "AWS", "React").
    for ent in doc.ents:
        if ent.label_ in RELEVANT_ENTITY_LABELS and not _is_noise_span(ent):
            candidates.append(ent.text)

    cleaned = []
    seen = set()
    for raw in candidates:
        phrase = _normalize_whitespace(raw)
        phrase = _strip_edge_punct(phrase)
        phrase = _strip_filler(phrase)
        phrase = _strip_edge_punct(phrase)

        if not phrase or len(phrase) < MIN_PHRASE_CHARS:
            continue

        word_count = len(phrase.split())
        if word_count > MAX_PHRASE_WORDS:
            continue

        if _is_generic_noise(phrase):
            continue

        key = phrase.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(phrase)

    return cleaned


def extract_skills_from_file(file_path) -> list:
    """Convenience wrapper: parse a PDF/DOCX file and extract skill phrases."""
    text = extract_text_from_file(file_path)
    return extract_skills_from_text(text)


# ---------------------------------------------------------------------------
# Self-test
# ---------------------------------------------------------------------------

SAMPLE_SYLLABUS_TEXT = """
This course provides an introduction to Python programming and covers the
fundamentals of data structures and algorithms. Students will study
object-oriented programming, relational database design using SQL, and the
basics of machine learning. The course also covers web development with
React and Node.js, version control using Git, and an overview of cloud
computing on AWS. By the end of the semester, students will have a strong
understanding of software engineering principles and agile methodologies.
"""


def test_extract_skills_from_text():
    """Sanity-check extraction against a sample syllabus paragraph."""
    skills = extract_skills_from_text(SAMPLE_SYLLABUS_TEXT)

    print("Extracted skill/topic phrases:")
    for skill in skills:
        print(f"  - {skill}")

    assert skills, "Expected at least one skill phrase to be extracted"
    assert not any(s.lower().startswith("introduction to") for s in skills), (
        "Filler phrase 'introduction to' was not stripped"
    )
    assert not any(s.lower().startswith("fundamentals of") for s in skills), (
        "Filler phrase 'fundamentals of' was not stripped"
    )

    print(f"\nSelf-test passed: {len(skills)} clean phrase(s) extracted.")
    return skills


if __name__ == "__main__":
    test_extract_skills_from_text()
